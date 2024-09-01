import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Asistente de Propuestas de Negocios", page_icon="💼", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Asistente de Propuestas de Negocios que ayuda a los usuarios a buscar datos de mercado, analizar tendencias, y evaluar competidores para generar una propuesta o plan de negocios detallado.

    ### Cómo usar la aplicación:

    1. Ingrese una idea de negocio o sector en el campo de texto.
    2. Haga clic en "Buscar información" para iniciar la búsqueda.
    3. Revise la propuesta de negocio generada, que incluye análisis de riesgos, estrategias de marketing, y proyecciones financieras.
    4. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **[Tu Nombre]**, [Fecha Actual]

    ### Cómo citar esta aplicación (formato APA):
    [Tu Apellido], [Inicial del Nombre]. (2024). *Asistente de Propuestas de Negocios* [Aplicación web]. [URL de tu aplicación]

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar propuestas de negocio. Siempre verifique la información con las fuentes originales para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Asistente de Propuestas de Negocios")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    def buscar_informacion(query):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} market data, trends, competitors",
            "num": 10
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generar_propuesta(titulo, snippet):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Datos del mercado: {snippet}\n\nGenera una propuesta de negocio detallada basada en esta información. Incluir análisis de riesgos, estrategias de marketing, y proyecciones financieras.\n\nPropuesta de negocio:",
            "max_tokens": 2048,
            "temperature": 0.2,
            "top_p": 0.9,
            "top_k": 50,
            "repetition_penalty": 1.1,
            "stop": ["Datos del mercado:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(idea, resultados):
        doc = Document()
        doc.add_heading('Propuesta de Negocio', 0)

        doc.add_heading('Idea de Negocio', level=1)
        doc.add_paragraph(idea)

        for resultado in resultados:
            doc.add_heading(resultado['title'], level=2)
            doc.add_paragraph(f"URL: {resultado['link']}")
            doc.add_paragraph(resultado['propuesta'])

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con las fuentes originales para un análisis más profundo.')

        return doc

    idea_negocio = st.text_input("Ingresa tu idea de negocio o sector:")

    if st.button("Buscar información"):
        if idea_negocio:
            with st.spinner("Buscando datos de mercado y generando propuesta..."):
                resultados_busqueda = buscar_informacion(idea_negocio)
                resultados = []

                for item in resultados_busqueda.get("organic", []):
                    titulo = item.get("title", "")
                    snippet = item.get("snippet", "")
                    link = item.get("link", "")

                    propuesta_negocio = generar_propuesta(titulo, snippet)

                    resultado = {
                        "title": titulo,
                        "link": link,
                        "propuesta": propuesta_negocio
                    }
                    resultados.append(resultado)

                # Mostrar los resultados
                st.subheader(f"Propuesta para la idea de negocio: {idea_negocio}")
                for resultado in resultados:
                    st.markdown(f"### [{resultado['title']}]({resultado['link']})")
                    st.markdown(resultado['propuesta'])
                    st.markdown("---")

                # Botón para descargar el documento
                doc = create_docx(idea_negocio, resultados)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar propuesta en DOCX",
                    data=buffer,
                    file_name=f"Propuesta_{idea_negocio.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, ingresa una idea de negocio o sector.")
